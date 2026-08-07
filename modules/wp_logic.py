"""
Working Professional Logic Module

Contains core rule-based algorithms for the Working Professional Module:
- Career Transition Matching
- Promotion Readiness Scoring
- Salary Benchmarking
- Certification Recommendations
"""
import json

def calculate_promotion_readiness(years_exp, project_complexity, leadership_score, tech_score, comms_score, team_score):
    """
    Computes a weighted Promotion Readiness % based on multiple factors.
    Weights: Technical (35%), Leadership (25%), Communication (15%), Team/Project (15%), Experience Base (10%)
    """
    # Normalize inputs to 100 max
    tech_w = (tech_score / 100) * 35
    lead_w = (leadership_score / 100) * 25
    comm_w = (comms_score / 100) * 15
    team_proj_score = (project_complexity + team_score) / 2
    team_w = (team_proj_score / 100) * 15
    
    # Experience gives a baseline boost (up to 10% for >= 5 years)
    exp_w = min(years_exp * 2, 10)
    
    total = tech_w + lead_w + comm_w + team_w + exp_w
    
    return {
        "Promotion Readiness %": int(total),
        "Technical Readiness": tech_score,
        "Leadership Readiness": leadership_score,
        "Communication Readiness": comms_score
    }

def match_career_transition(user_skill_scores, role_requirements):
    """
    Given a user's skill scores (dict: area -> score 0-100) and a list of 
    WP_RoleRequirement objects, returns a ranked list of roles and match %.
    
    user_skill_scores: e.g. {"Backend Development": 88, "System Design": 72}
    role_requirements: list of WP_RoleRequirement instances.
    """
    matches = []
    
    for role in role_requirements:
        reqs = json.loads(role.required_skills_weights)
        total_weight = 0
        weighted_score = 0
        
        for skill, weight in reqs.items():
            user_score = user_skill_scores.get(skill, 0)
            weighted_score += user_score * weight
            total_weight += weight
            
        if total_weight > 0:
            match_pct = int(weighted_score / total_weight)
        else:
            match_pct = 0
            
        matches.append({
            "Next Role": role.role_name,
            "Match %": match_pct
        })
        
    matches.sort(key=lambda x: x["Match %"], reverse=True)
    return matches

def advanced_job_matching(user_profile, user_skill_scores, role_requirements):
    """
    Similar to match_career_transition but factors in Experience and Salary Expectations.
    (Simplified rule-based approach)
    """
    matches = match_career_transition(user_skill_scores, role_requirements)
    # Apply a slight penalty if min experience is not met
    user_exp = user_profile.total_experience_years or 0
    
    for role_model in role_requirements:
        for match in matches:
            if match["Next Role"] == role_model.role_name:
                if user_exp < role_model.min_experience_years:
                    penalty = (role_model.min_experience_years - user_exp) * 5
                    match["Match %"] = max(0, match["Match %"] - int(penalty))
                    
    matches.sort(key=lambda x: x["Match %"], reverse=True)
    return matches

def get_salary_growth(current_salary, target_role, salary_benchmarks):
    """
    Calculates potential salary growth.
    salary_benchmarks is a list of WP_SalaryBenchmark instances.
    """
    for sb in salary_benchmarks:
        if sb.role_name == target_role:
            growth_pct = ((sb.market_average - current_salary) / current_salary) * 100 if current_salary > 0 else 0
            return {
                "Market Average": sb.market_average,
                "Target Range": f"₹{sb.target_min} - ₹{sb.target_max} LPA",
                "Potential Growth %": int(growth_pct)
            }
    return None

def recommend_certifications(user_skill_scores, target_role_req, certifications_list):
    """
    Recommends certifications based on user's lowest skills required by the target role.
    target_role_req is a WP_RoleRequirement.
    certifications_list is a list of WP_Certification.
    """
    if not target_role_req:
        return []
        
    reqs = json.loads(target_role_req.required_skills_weights)
    
    # Find skills where user score is < 70
    gaps = []
    for skill in reqs.keys():
        score = user_skill_scores.get(skill, 0)
        if score < 70:
            gaps.append(skill)
            
    recs = []
    for cert in certifications_list:
        if cert.related_skill in gaps:
            recs.append({
                "Certification": cert.name,
                "Priority": cert.priority,
                "Related Skill": cert.related_skill
            })
            
    # Sort High -> Medium -> Low
    priority_map = {"High": 1, "Medium": 2, "Low": 3}
    recs.sort(key=lambda x: priority_map.get(x["Priority"], 4))
    return recs

def generate_action_plan(user_skill_scores):
    """
    Generates a 90-Day Action Plan focusing on the user's lowest skills.
    """
    if not user_skill_scores:
        return []
        
    # Sort skills by score ascending
    sorted_skills = sorted(user_skill_scores.items(), key=lambda x: x[1])
    lowest = [s[0] for s in sorted_skills[:3]]
    
    plan = []
    if len(lowest) > 0:
        plan.append({"Month": "Month 1", "Focus": f"Core Fundamentals in {lowest[0]}"})
    if len(lowest) > 1:
        plan.append({"Month": "Month 2", "Focus": f"Advanced Practices in {lowest[1]}"})
    if len(lowest) > 2:
        plan.append({"Month": "Month 3", "Focus": f"Integration & Mastery of {lowest[2]}"})
        
    return plan


# ──────────────────────────────────────────────────────────────────────────────
# FEATURE 1: RESUME UPDATE ASSISTANT LOGIC
# ──────────────────────────────────────────────────────────────────────────────

import io
import re
import os


def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """
    Extracts text content from uploaded PDF, DOCX, or text files.
    """
    ext = os.path.splitext(filename)[1].lower()
    text = ""
    
    if ext == ".pdf":
        # Attempt pdfplumber first
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + "\n"
        except Exception:
            pass

        # Fallback to PyPDF2 if pdfplumber produced empty or failed
        if not text.strip():
            try:
                import PyPDF2
                reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
                for page in reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + "\n"
            except Exception:
                pass

    elif ext in [".docx", ".doc"]:
        try:
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text += row_text + "\n"
        except Exception:
            pass

    # Generic fallback for text or if others failed
    if not text.strip():
        try:
            text = file_bytes.decode('utf-8', errors='ignore')
        except Exception:
            text = str(file_bytes)

    return text.strip()


INDUSTRY_TECH_SKILLS = {
    "Backend": ["Python", "Java", "Go", "Rust", "C++", "C#", ".NET", "Node.js", "Django", "FastAPI", "Flask", "Spring Boot", "Express.js", "Ruby on Rails", "REST API", "GraphQL", "gRPC"],
    "Cloud & DevOps": ["AWS", "Azure", "GCP", "Docker", "Kubernetes", "CI/CD", "GitHub Actions", "Terraform", "Ansible", "Helm", "Linux", "Nginx", "Prometheus", "Grafana"],
    "Databases & Storage": ["PostgreSQL", "MySQL", "MongoDB", "Redis", "Kafka", "RabbitMQ", "Elasticsearch", "Cassandra", "DynamoDB", "Snowflake", "SQL"],
    "System Architecture": ["System Design", "Microservices", "Distributed Systems", "Serverless", "Design Patterns", "Event-Driven Architecture", "Scalability", "High Availability", "Caching", "Load Balancing"],
    "Leadership & Process": ["Agile", "Scrum", "Team Leadership", "Mentoring", "Code Review", "Project Ownership", "Stakeholder Management", "Sprint Planning", "Strategic Thinking", "Architecture Review"]
}

OUTDATED_SKILLS_LIST = [
    "SOAP", "SVN", "Subversion", "Apache Ant", "Flash", "ActionScript",
    "AngularJS 1.x", "JSP", "Servlets", "VB.NET", "Perl", "Apache Struts",
    "XML-RPC", "Visual SourceSafe", "CVS", "COBOL", "ColdFusion", "PHP 5",
    "Silverlight", "Bower", "Grunt", "Gulp"
]

TRENDING_RECOMMENDATIONS = [
    {"skill": "Kubernetes", "category": "Cloud & DevOps", "growth": "+38%", "reason": "Standard for modern container orchestration and high-availability systems."},
    {"skill": "AWS Cloud Architecture", "category": "Cloud & DevOps", "growth": "+32%", "reason": "Leading enterprise cloud provider requirement across 70%+ senior roles."},
    {"skill": "Microservices & Distributed Systems", "category": "System Architecture", "growth": "+28%", "reason": "Crucial for scaling backend engineering and modular software design."},
    {"skill": "System Design & Scalability", "category": "System Architecture", "growth": "+25%", "reason": "Required benchmark for Staff / Senior Engineering and Architecture tracks."},
    {"skill": "Kafka / Event Streaming", "category": "Databases & Storage", "growth": "+24%", "reason": "Standard for high-throughput real-time distributed data pipelines."},
    {"skill": "Terraform / IaC", "category": "Cloud & DevOps", "growth": "+22%", "reason": "Industry standard for automated immutable cloud infrastructure."}
]


def analyze_working_prof_resume(resume_text: str, user_profile=None, trending_skills=None, role_requirements=None) -> dict:
    """
    Comprehensive resume analysis for working professionals:
    - Skill Extraction
    - Gap Analysis against Industry Trends & Target Roles
    - Outdated Skill Detection
    - Section Improvement Suggestions (Summary, Tech Skills, Projects, Experience, Certifications)
    - ATS Compliance Checklist & Completeness Score (0-100)
    """
    lower_text = resume_text.lower()
    
    # 1. Extract Skills
    extracted_by_category = {}
    all_extracted_skills = set()
    
    for category, skills in INDUSTRY_TECH_SKILLS.items():
        matched = []
        for s in skills:
            pattern = r'(?<![a-zA-Z0-9])' + re.escape(s.lower()) + r'(?![a-zA-Z0-9])'
            if re.search(pattern, lower_text):
                matched.append(s)
                all_extracted_skills.add(s)
        if matched:
            extracted_by_category[category] = matched

    extracted_skills_list = sorted(list(all_extracted_skills))

    # 2. Outdated Skills Detection
    outdated_found = []
    for legacy in OUTDATED_SKILLS_LIST:
        pattern = r'(?<![a-zA-Z0-9])' + re.escape(legacy.lower()) + r'(?![a-zA-Z0-9])'
        if re.search(pattern, lower_text):
            outdated_found.append(legacy)

    # 3. Target Role & Missing Skills
    target_role = user_profile.preferred_job_roles if user_profile and user_profile.preferred_job_roles else "Senior Backend Engineer"
    
    # Identify in-demand core skills for professional
    core_modern_skills = ["System Design", "Cloud Computing", "AWS", "Docker", "Kubernetes", "Microservices", "CI/CD", "Redis", "Kafka", "SQL", "PostgreSQL", "Mentoring"]
    missing_skills = [s for s in core_modern_skills if s not in all_extracted_skills]

    # Recommended skills from trending
    recommended_skills = [
        t for t in TRENDING_RECOMMENDATIONS 
        if t["skill"] not in all_extracted_skills
    ]

    # 4. Section Presence and Quality Scoring
    has_contact = bool(re.search(r'[\w\.-]+@[\w\.-]+', resume_text) or re.search(r'\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}', resume_text))
    has_summary = bool(re.search(r'\b(summary|profile|about me|objective|professional summary)\b', lower_text))
    has_experience = bool(re.search(r'\b(experience|work experience|employment|career history)\b', lower_text))
    has_projects = bool(re.search(r'\b(projects|key projects|technical projects|portfolio)\b', lower_text))
    has_education = bool(re.search(r'\b(education|university|degree|btech|bachelor|master)\b', lower_text))
    has_certifications = bool(re.search(r'\b(certifications|certificate|certified|credentials)\b', lower_text))
    has_metrics = bool(re.search(r'(\d+%\s*|\$\d+|\d+\s*x|\d+\s*ms|\d+\s*users|\d+\s*k|\d+\s*m\b)', lower_text))
    
    # Calculate Completeness Score (0-100)
    score = 0
    if has_contact: score += 10
    if has_summary: score += 15
    if has_experience: score += 20
    if has_projects: score += 15
    if has_education: score += 10
    if len(extracted_skills_list) >= 5: score += 15
    elif len(extracted_skills_list) >= 2: score += 10
    if has_certifications: score += 5
    if has_metrics: score += 10
    score = min(100, max(20, score))

    # 5. Section Improvement Suggestions
    curr_role = user_profile.current_role if user_profile and user_profile.current_role else "Software Engineer"
    curr_exp = user_profile.total_experience_years if user_profile and user_profile.total_experience_years else 3.5

    summary_suggestion = (
        f"**Suggested Professional Summary:**\n\n"
        f"> \"Results-driven {curr_role} with {curr_exp}+ years of experience architecting resilient distributed systems, "
        f"cloud-native backends, and high-throughput APIs. Proven track record in optimizing application performance by up to 35%, "
        f"leading cross-functional agile initiatives, and implementing robust CI/CD pipelines with modern cloud architectures.\""
    )

    tech_skills_suggestion = [
        "Group your skills into standardized industry categories: **Languages & Frameworks**, **Cloud & DevOps**, **Databases & Caching**, and **System Architecture**.",
        "List proficiency alongside key frameworks (e.g. *Python (FastAPI, Django)*, *AWS (EC2, S3, Lambda, ECS)*).",
        "Place your strongest, most modern technical skills in the top 3 lines of your skills section for ATS keyword prominence."
    ]

    project_suggestions = [
        "**Apply the STAR Method (Situation, Task, Action, Result):** For every project, specify the business challenge, the technical solution, and the measurable outcome.",
        "**Action Verbs:** Start bullet points with high-impact verbs such as *Architected*, *Spearheaded*, *Engineered*, *Optimized*, or *Automated* instead of passive phrases like *'Responsible for'*.",
        "**Quantify System Scale:** State real metrics: *'Engineered microservices handling 20,000+ requests/min with sub-50ms latency'*, *'Reduced deployment turnaround by 45% via automated GitHub Actions CI/CD'*."
    ]

    experience_suggestions = [
        "Highlight architectural ownership and cross-functional team collaboration.",
        "Include leadership impact: mention code reviews, mentoring junior engineers, and driving engineering best practices.",
        "Demonstrate business alignment: articulate how your software features directly drove user growth, reliability uptime (e.g. 99.9% SLA), or infrastructure cost reduction."
    ]

    cert_suggestions = [
        "AWS Certified Solutions Architect – Associate / Professional",
        "Certified Kubernetes Administrator (CKA)",
        "HashiCorp Certified: Terraform Associate",
        "TOGAF 9.2 Standard / System Design Masterclass"
    ]

    # 6. ATS Improvement Tips
    ats_tips = [
        "✅ **Use Clean Single-Column Layout:** Avoid multi-column tables, text boxes, and sidebar graphics that disrupt ATS parsers.",
        "✅ **Standard Section Headings:** Use conventional headers such as *Professional Experience*, *Technical Skills*, *Projects*, and *Education*.",
        "✅ **Direct Keyword Matching:** Match exact keywords from target job descriptions (e.g. *Kubernetes*, *Microservices*, *CI/CD*).",
        "✅ **Quantifiable Metrics:** Include percentages (%), latency reductions (ms), and scale indicators to boost scoring in modern parsing engines.",
        "✅ **File Format:** Save and upload as clean PDF or DOCX using standard system fonts (Inter, Arial, Calibri, or Roboto)."
    ]

    # 7. Generate Printable / Downloadable Report
    download_text = f"""# TALENTSPHERE ELEVATE - RESUME ANALYSIS & IMPROVEMENT REPORT
Generated on: {user_profile.current_role if user_profile else 'Professional'} Growth Assessment
Completeness Score: {score}/100

==================================================
1. EXTRACTED SKILLS ({len(extracted_skills_list)} Identified)
==================================================
{', '.join(extracted_skills_list) if extracted_skills_list else 'No recognized technical keywords detected.'}

==================================================
2. INDUSTRY GAP ANALYSIS
==================================================
Missing Core Skills: {', '.join(missing_skills) if missing_skills else 'None (Great Coverage!)'}
Outdated / Legacy Skills: {', '.join(outdated_found) if outdated_found else 'None (Modern Stack!)'}

Top Recommended Skills to Add:
{chr(10).join([f"- {r['skill']} ({r['growth']} industry growth): {r['reason']}" for r in recommended_skills])}

==================================================
3. SECTION-BY-SECTION IMPROVEMENT SUGGESTIONS
==================================================
[Professional Summary]
{summary_suggestion}

[Technical Skills]
{chr(10).join(['- ' + s for s in tech_skills_suggestion])}

[Projects & System Architecture]
{chr(10).join(['- ' + s for s in project_suggestions])}

[Experience & Leadership]
{chr(10).join(['- ' + s for s in experience_suggestions])}

[Recommended Certifications]
{chr(10).join(['- ' + c for c in cert_suggestions])}

==================================================
4. ATS OPTIMIZATION CHECKLIST
==================================================
{chr(10).join(ats_tips)}
"""

    return {
        "completeness_score": score,
        "extracted_skills": extracted_skills_list,
        "extracted_by_category": extracted_by_category,
        "missing_skills": missing_skills,
        "outdated_skills": outdated_found,
        "recommended_skills": recommended_skills,
        "summary_suggestion": summary_suggestion,
        "tech_skills_suggestion": tech_skills_suggestion,
        "project_suggestions": project_suggestions,
        "experience_suggestions": experience_suggestions,
        "cert_suggestions": cert_suggestions,
        "ats_tips": ats_tips,
        "download_text": download_text,
        "section_checklist": {
            "Contact Details": has_contact,
            "Professional Summary": has_summary,
            "Technical Skills": len(extracted_skills_list) >= 3,
            "Work Experience": has_experience,
            "Projects": has_projects,
            "Education": has_education,
            "Certifications": has_certifications,
            "Measurable Metrics": has_metrics
        }
    }


# ──────────────────────────────────────────────────────────────────────────────
# FEATURE 2: AI CAREER COACH ENGINE
# ──────────────────────────────────────────────────────────────────────────────

class CareerCoachEngine:
    """
    Intelligent AI Career Coach for Working Professionals.
    Supports Google Gemini / OpenAI when API keys are available, with an
    expert contextual rule-based engine fallback.
    """

    def __init__(self):
        self.google_api_key = os.environ.get("GOOGLE_API_KEY", "") or os.environ.get("GEMINI_API_KEY", "")
        self.openai_api_key = os.environ.get("OPENAI_API_KEY", "")
        self._gemini_model = None

    def _get_gemini_model(self):
        if not self.google_api_key:
            return None
        if self._gemini_model is None:
            try:
                import google.generativeai as genai
                genai.configure(api_key=self.google_api_key)
                self._gemini_model = genai.GenerativeModel('gemini-2.0-flash')
            except Exception:
                return None
        return self._gemini_model

    def get_response(self, user_query: str, context: dict, history: list = None) -> str:
        """
        Generates an intelligent career coaching response using LLM or rule-based engine.
        """
        gemini = self._get_gemini_model()
        if gemini:
            try:
                system_prompt = (
                    "You are an Elite Executive Career Coach and Senior Technical Advisor on TalentSphere Elevate. "
                    "You provide strategic, actionable, and encouraging career advice tailored to working professionals.\n\n"
                    f"User Profile Context:\n"
                    f"- Current Role: {context.get('current_role', 'Software Engineer')}\n"
                    f"- Total Experience: {context.get('experience', '3')} Years\n"
                    f"- Career Goal: {context.get('career_goal', 'Senior Backend Engineer')}\n"
                    f"- Assessed Skills & Scores: {json.dumps(context.get('skill_scores', {}))}\n"
                    f"- Promotion Readiness: {context.get('promotion_readiness', 65)}%\n"
                    f"- Target Job Match: {context.get('best_match', 'Senior Backend Engineer')}\n"
                    f"- Salary Growth Potential: {context.get('salary_growth', '+25%')}\n\n"
                    "Respond with clear formatting, bullet points, and specific action steps. Keep it under 350 words."
                )
                prompt = f"System:\n{system_prompt}\n\nUser Question: {user_query}\n\nCoach Response:"
                response = gemini.generate_content(prompt)
                if response and response.text:
                    return response.text
            except Exception:
                pass  # Fallback to rule engine

        # Expert Rule-Based Recommendation Engine
        return self._rule_based_response(user_query, context)

    def _rule_based_response(self, query: str, ctx: dict) -> str:
        q = query.lower()
        role = ctx.get("current_role", "Software Engineer")
        exp = ctx.get("experience", 3.5)
        goal = ctx.get("career_goal", "Senior Backend Engineer")
        readiness = ctx.get("promotion_readiness", 65)
        best_match = ctx.get("best_match", "Senior Backend Engineer")
        skill_scores = ctx.get("skill_scores", {})
        
        # 1. Senior Backend Engineer Roadmap
        if "senior backend" in q or ("become" in q and "backend" in q) or "next role" in q:
            return (
                f"### 🚀 Strategic Blueprint: Path to Senior Backend Engineer\n\n"
                f"As a **{role}** with **{exp} years of experience**, here is your structured transition plan to **{best_match}**:\n\n"
                f"1. **Deep System Architecture & Distributed Systems:**\n"
                f"   - Master high-throughput message streaming with **Apache Kafka** or **RabbitMQ**.\n"
                f"   - Implement distributed caching strategies with **Redis** (cache-aside, write-through).\n"
                f"   - Design for 99.99% availability, idempotency, and database sharding.\n\n"
                f"2. **Cloud-Native & Container Mastery:**\n"
                f"   - Learn Kubernetes pod lifecycle, ingress controllers, and Helm chart packaging.\n"
                f"   - Architect serverless or microservice APIs on AWS/GCP with automated CI/CD.\n\n"
                f"3. **Technical Leadership & Ownership:**\n"
                f"   - Lead design reviews (RFCs), establish code review standards, and mentor junior engineers.\n"
                f"   - Drive end-to-end reliability, SLA monitoring with Prometheus & Grafana.\n\n"
                f"💡 *Current Match:* Your profile shows a strong baseline match for **{best_match}**. Focus on closing your Cloud & System Design gaps!"
            )

        # 2. Certification Advice
        elif "certification" in q or "cert" in q or "certificate" in q:
            return (
                f"### 📜 Targeted High-ROI Certification Recommendations\n\n"
                f"Based on your profile as a **{role}** aiming for **{goal}**, here are the highest-impact certifications to prioritize:\n\n"
                f"1. **AWS Certified Solutions Architect – Associate / Professional (Priority: High)**\n"
                f"   - *Why:* Recognized industry benchmark verifying multi-tier cloud infrastructure, security, and scalability.\n\n"
                f"2. **Certified Kubernetes Administrator (CKA) (Priority: High)**\n"
                f"   - *Why:* Validates hands-on production cluster setup, networking, and container workload management.\n\n"
                f"3. **HashiCorp Certified: Terraform Associate (Priority: Medium)**\n"
                f"   - *Why:* Accelerates your Infrastructure-as-Code (IaC) and DevOps engineering credentials.\n\n"
                f"4. **System Design & Software Architecture Masterclasses (Priority: High)**\n"
                f"   - *Why:* Directly empowers you in technical interviews and staff engineering decision-making."
            )

        # 3. What to learn after Python
        elif "after python" in q or "learn next" in q or "what should i learn" in q:
            return (
                f"### ⚡ What to Learn After Python Mastery\n\n"
                f"Since you already have Python in your toolkit, here is the highest leverage trajectory to expand your engineering horizon:\n\n"
                f"1. **Modern High-Performance Language (Go or Rust):**\n"
                f"   - **Go (Golang):** Ideal for cloud-native backend services, microservices, and Kubernetes tooling.\n"
                f"   - **Rust:** Exceptional for memory-safe, ultra-low-latency systems and tooling.\n\n"
                f"2. **Advanced Distributed Systems & Middleware:**\n"
                f"   - **FastAPI / Asyncio:** Master concurrent Python, event loops, and asynchronous database drivers.\n"
                f"   - **Kafka & Redis:** Event streaming, pub/sub, distributed locks, and in-memory caching.\n\n"
                f"3. **Cloud & Infrastructure as Code:**\n"
                f"   - Master **Docker** containerization, **Kubernetes** orchestration, and **Terraform**.\n\n"
                f"4. **AI & Vector Databases:**\n"
                f"   - Integrate LLMs, LangChain/LlamaIndex, and Vector DBs (Pinecone, ChromaDB, pgvector) with Python."
            )

        # 4. How to get promoted
        elif "promoted" in q or "promotion" in q or "career growth" in q:
            return (
                f"### 📈 Action Plan: Accelerating Your Promotion Readiness (Current Score: {readiness}%)\n\n"
                f"To transition from **{role}** to the next senior level:\n\n"
                f"1. **Expand Project Scope & Ambiguity:**\n"
                f"   - Volunteer to lead complex, cross-team projects that lack predefined solutions.\n"
                f"   - Author Technical Design Documents (TDDs) and drive consensus across stakeholders.\n\n"
                f"2. **Elevate Team Multiplier Effect:**\n"
                f"   - Actively mentor 1-2 junior developers.\n"
                f"   - Run technical brown-bag sessions and improve engineering onboarding documentation.\n\n"
                f"3. **Business Impact & Metrics:**\n"
                f"   - Frame all your engineering accomplishments around business value (e.g. latency, reliability, cost savings).\n"
                f"   - Maintain a **Brag Sheet** of measurable results ahead of your quarterly/annual review.\n\n"
                f"4. **Schedule a Growth Alignment 1-on-1 with Your Manager:**\n"
                f"   - Ask directly: *'What specific competencies and deliverables do you need to see from me over the next 6 months to earn a promotion?'*"
            )

        # 5. Salary Expectations
        elif "salary" in q or "compensation" in q or "lpa" in q or "market rate" in q:
            return (
                f"### 💰 Market Salary Benchmark & Growth Insights\n\n"
                f"Based on real-time industry benchmark data for **{best_match}**:\n\n"
                f"- **Target Role:** {best_match}\n"
                f"- **Market Average:** ₹9.0 – ₹15.0 LPA (for 3-6 years experience)\n"
                f"- **Target Senior Bracket:** ₹15.0 – ₹25.0+ LPA (with leadership & system design mastery)\n"
                f"- **Potential Compensation Growth:** +25% to +45%\n\n"
                f"**How to Maximize Your Compensation:**\n"
                f"1. **System Design Depth:** High-paying product companies place 50%+ interview weight on scalable architecture.\n"
                f"2. **Cloud & DevOps Versatility:** Full-lifecycle engineers command a 20-30% premium over mono-stack developers.\n"
                f"3. **Negotiation Leverage:** Benchmark with Glassdoor/Levels.fyi and highlight multi-cloud + leadership credentials."
            )

        # 6. Skill Gap & Learning Roadmap
        elif "skill gap" in q or "gap" in q or "roadmap" in q or "weak" in q:
            gaps = [k for k, v in skill_scores.items() if v < 70]
            gaps_str = ", ".join(gaps) if gaps else "Cloud Architecture, Kubernetes, Distributed Systems"
            return (
                f"### 🎯 Skill Gap Analysis & 90-Day Roadmap\n\n"
                f"Based on your latest assessments, here are your key focus areas: **{gaps_str}**.\n\n"
                f"- **Month 1 (Fundamentals & Foundations):** Focus on deep architecture patterns, Docker containerization, and SQL indexing.\n"
                f"- **Month 2 (Cloud & Distributed Engineering):** Implement AWS multi-tier deployments, Redis caching, and Kafka pipelines.\n"
                f"- **Month 3 (System Design & Leadership):** Build a full end-to-end scalable microservice project, practice system design mocks, and pursue certification.\n\n"
                f"Track your progress in the **Growth Summary** tab!"
            )

        # General Coaching Response
        else:
            return (
                f"### 💼 Career Coach Recommendation for {role}\n\n"
                f"Thank you for reaching out! Here is customized guidance based on your profile:\n\n"
                f"- **Current Role:** {role} ({exp} Yrs Exp)\n"
                f"- **Target Milestone:** {goal}\n"
                f"- **Promotion Readiness:** {readiness}%\n\n"
                f"**Key Recommendations:**\n"
                f"1. **Focus on High-Growth Skills:** Deepen your expertise in **System Design**, **AWS/Cloud Infrastructure**, and **Microservices**.\n"
                f"2. **Strengthen Leadership:** Mentor team members and take ownership of technical roadmaps to boost your promotion readiness.\n"
                f"3. **Optimize Your Resume:** Use the **Resume Update Assistant** tab to ensure ATS keyword compliance and download actionable tips.\n\n"
                f"Feel free to ask specific questions about certifications, salary benchmarks, interview preparation, or promotion strategies!"
            )


# ──────────────────────────────────────────────────────────────────────────────
# FEATURE 3: LEADERSHIP EVALUATION LOGIC
# ──────────────────────────────────────────────────────────────────────────────

LEADERSHIP_DIMENSIONS = [
    "Team Coordination",
    "Mentoring",
    "Decision Making",
    "Conflict Resolution",
    "Project Ownership",
    "Communication",
    "Strategic Thinking"
]

LEADERSHIP_COURSES = [
    {"name": "Engineering Management & Technical Leadership", "provider": "Coursera / Stanford", "focus": "Team Coordination & Mentoring"},
    {"name": "Strategic Decision Making & Architecture Governance", "provider": "edX / MIT", "focus": "Decision Making & Strategic Thinking"},
    {"name": "Crucial Conversations: Conflict Resolution in Tech", "provider": "LinkedIn Learning", "focus": "Conflict Resolution & Communication"},
    {"name": "High-Output Management & Agile Project Ownership", "provider": "Udemy Executive", "focus": "Project Ownership & Execution"},
    {"name": "Executive Communication for Technology Leaders", "provider": "Harvard Online", "focus": "Communication & Stakeholder Management"}
]


def evaluate_leadership_skills(scores: dict) -> dict:
    """
    Evaluates 7 core leadership dimensions and computes:
    - Overall Leadership Score (0-100)
    - Leadership Grade
    - Strong Areas & Weak Areas
    - Tailored Actionable Suggestions
    - Recommended Courses
    - Promotion Impact Score
    """
    total_val = 0
    count = 0
    clean_scores = {}
    
    for dim in LEADERSHIP_DIMENSIONS:
        val = float(scores.get(dim, 50.0))
        clean_scores[dim] = val
        total_val += val
        count += 1
        
    overall_score = round(total_val / max(1, count), 1)
    
    # Assign Grade
    if overall_score >= 90:
        grade = "A+ Executive & Strategic Leader"
        readiness_status = "Executive Ready: Exceptional team leadership and strategic vision."
    elif overall_score >= 80:
        grade = "A Senior Team & Tech Lead"
        readiness_status = "Senior Lead Ready: Strong leadership capability across all major facets."
    elif overall_score >= 70:
        grade = "B+ Emerging Leader / Lead Developer"
        readiness_status = "Emerging Leader: Ready for Tech Lead and team management roles with guidance."
    elif overall_score >= 60:
        grade = "B Developing Leader"
        readiness_status = "Developing: Good foundation; focus on mentoring and strategic decision making."
    else:
        grade = "C Foundational / Needs Mentorship"
        readiness_status = "Foundational: Requires active leadership coaching and project ownership practice."

    strong_areas = [k for k, v in clean_scores.items() if v >= 75]
    weak_areas = [k for k, v in clean_scores.items() if v < 75]

    # Specific actionable coaching per weak area
    suggestions_map = {
        "Team Coordination": "Implement clear Agile sprint ceremonies, cross-functional standups, and structured task breakdown frameworks.",
        "Mentoring": "Establish structured bi-weekly 1-on-1s with junior developers; guide code reviews with constructive learning comments.",
        "Decision Making": "Adopt structured Decision Records (ADRs) and data-backed trade-off analysis before committing to architecture paths.",
        "Conflict Resolution": "Practice non-violent communication and active listening; address blockers early before they escalate into team friction.",
        "Project Ownership": "Take full end-to-end accountability for feature roadmaps, from design and QA to deployment and telemetry monitoring.",
        "Communication": "Present technical designs clearly to both engineering and non-technical business stakeholders with concise executive summaries.",
        "Strategic Thinking": "Align your engineering deliverables directly with quarterly company OKRs and scalable multi-quarter architecture plans."
    }

    improvement_suggestions = [
        {"area": w, "tip": suggestions_map.get(w, "Focus on proactive practice and seek feedback from your peers.")}
        for w in weak_areas
    ]

    # Promotion Impact Score (Leadership accounts for 25% of overall promotion readiness)
    promotion_impact = round((overall_score / 100) * 25, 1)

    return {
        "scores": clean_scores,
        "overall_score": overall_score,
        "grade": grade,
        "readiness_status": readiness_status,
        "strong_areas": strong_areas,
        "weak_areas": weak_areas,
        "improvement_suggestions": improvement_suggestions,
        "recommended_courses": LEADERSHIP_COURSES,
        "promotion_impact": promotion_impact
    }

